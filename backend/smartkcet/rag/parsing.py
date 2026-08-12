"""Document parsing and OCR helpers.

These utilities were lifted verbatim from ``backend/app.py`` and grouped
here so the upload pipeline can be evolved (per-subject isolation, OCR
warning aggregation per design.md §5) without touching unrelated modules.
"""

from __future__ import annotations

import io
import logging
from typing import List

import cv2
import fitz
import numpy as np
import pytesseract
from PIL import Image
from docx import Document as DocxDocument

logger = logging.getLogger("smartkcet.rag.parsing")


def preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """Apply denoising and adaptive thresholding for better OCR accuracy."""

    try:
        img_np = np.array(img.convert("RGB"))
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        gray = cv2.fastNlMeansDenoising(gray, h=10)
        thresh = cv2.adaptiveThreshold(
            gray,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            31,
            10,
        )
        return Image.fromarray(thresh)
    except Exception as exc:  # pragma: no cover - defensive logging
        print(f"OCR preprocessing failed: {exc}")
        return img


def _process_groq_vision(img_str: str, page_num: int) -> str:
    from ..rag.groq_client import get_groq_client
    try:
        client = get_groq_client()
        completion = client.chat.completions.create(
            model="llama-3.2-90b-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all the text from this image accurately. Do not add any extra commentary or markdown formatting, just the extracted text. IMPORTANT: Do NOT extract or include any questions that rely on diagrams, figures, or images to be solved. Do NOT extract any irrelevant questions. Only extract proper academic questions related to the chapter."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img_str}"}}
                    ]
                }
            ],
            temperature=0.0,
        )
        return completion.choices[0].message.content or ""
    except Exception as groq_exc:
        logger.warning(f"Groq Vision failed on page {page_num+1}: {groq_exc}")
        return ""

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF; fall back to OCR for pages with little text."""

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    logger.info(
        "PDF opened: %d page(s), metadata=%r",
        len(doc),
        doc.metadata.get("title", ""),
    )
    
    import concurrent.futures
    import base64
    
    pages_texts = ["" for _ in range(len(doc))]
    futures = []

    with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
        for page_num, page in enumerate(doc):
            text = page.get_text().strip()
            if len(text) > 50:
                logger.debug(
                    "Page %d: extracted %d chars via text layer (direct extraction)",
                    page_num + 1,
                    len(text),
                )
                pages_texts[page_num] = text
                continue

            # Text layer is empty or too short — attempt OCR fallback
            base_text = text if text else ""
            if text:
                logger.info(
                    "Page %d: text layer has only %d chars (below 50-char threshold), "
                    "keeping it and also attempting OCR",
                    page_num + 1,
                    len(text),
                )
            else:
                logger.info(
                    "Page %d: text layer is empty, attempting OCR fallback",
                    page_num + 1,
                )
                
            try:
                pix = page.get_pixmap(dpi=300)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                img = preprocess_for_ocr(img)
                
                buffered = io.BytesIO()
                img.save(buffered, format="JPEG")
                img_str = base64.b64encode(buffered.getvalue()).decode()
                
                # Submit API call to thread pool
                future = executor.submit(_process_groq_vision, img_str, page_num)
                futures.append((page_num, base_text, future))
                
            except Exception as exc:
                logger.warning(
                    "Page %d: pixmap/OCR pipeline failed: %s",
                    page_num + 1,
                    exc,
                )
                pages_texts[page_num] = base_text

        # Wait for all futures
        for page_num, base_text, future in futures:
            try:
                ocr_text = future.result()
                if ocr_text.strip():
                    logger.info(
                        "Page %d: Groq Vision OCR produced %d chars",
                        page_num + 1,
                        len(ocr_text.strip()),
                    )
                    pages_texts[page_num] = base_text + "\n" + ocr_text if base_text else ocr_text
                else:
                    logger.warning(
                        "Page %d: Groq Vision OCR returned empty text (truly unreadable page)",
                        page_num + 1,
                    )
                    pages_texts[page_num] = base_text
            except Exception as e:
                logger.error("Page %d OCR Future failed: %s", page_num + 1, e)
                pages_texts[page_num] = base_text

    total_text = "\n".join(t for t in pages_texts if t.strip())
    logger.info(
        "PDF extraction complete: %d page(s) yielded text, total %d chars",
        len([t for t in pages_texts if t.strip()]),
        len(total_text),
    )
    return total_text


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    logger.info("DOCX extraction: %d chars from %d paragraphs", len(text), len(doc.paragraphs))
    return text


def extract_text_from_txt(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    logger.info("TXT extraction: %d chars", len(text))
    return text


def chunk_text(text: str, size: int = 400, overlap: int = 80) -> List[str]:
    """Split ``text`` into overlapping word-windows ready for embedding."""

    words = text.split()
    chunks: List[str] = []
    i = 0
    while i < len(words):
        chunks.append(" ".join(words[i : i + size]))
        i += size - overlap
    return [c for c in chunks if len(c.strip()) > 30]
